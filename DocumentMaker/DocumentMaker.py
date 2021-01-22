from docx import Document
import os
#from docx.shared import Inches

#def PadString(sTarget):

def get_current_date_time():
    dtCurrent = datetime.now()
    return dtCurrent.strftime('%d/%m/%Y %H:%M:%S')

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

document.save('demo.docx')

print(os.path.exists(os.path.join(os.getcwd(), 'demo.docx')))

TITLE_PAD_LENGTH = 20
HOT_FIX_DESCRRIPTION = ' (Hot Fix)'
ReleaseUser = os.getenv('RELEASE_USER')
Result = [] #Result := TStringList.Create;
  #try
   # try
      #for i := 0 to FReleasedProducts.Count - 1 do
      #begin
      #  aCurProject   :=  FReleasedProducts.Objects[i] as TReleasedProduct;
      #  aProjectsName :=  aProjectsName + ' ' + aCurProject.Name;
      #end;

      #ResStrings := TstringList.Create;

Result.append('')
Result.append('Patch Release Document.')#Result.Add('Patch Release Document.');
Result.append('---------------------------------')#Result.Add(' ---------------------------------');
#//      Result.Add('');
Result.append('')#      Result.Add('');
Result.append('Patch Number'.ljust(TITLE_PAD_LENGTH))#      Result.Add(PadChar('Patch Number', TITLE_PAD_LENGTH, ' ')   +  FVersionNum);
Result.append(f"{'Date Packed'.ljust(TITLE_PAD_LENGTH) {get_current_date_time()}"})#      Result.Add(PadChar('Date Packed', TITLE_PAD_LENGTH, ' ') +  GetStringFromIntDate(FReleaseDate) + ' ' + GetStringFromIntTime(FReleaseTime));
Result.append(f"{'Developer Name'.ljust(TITLE_PAD_LENGTH)} {ReleaseUser}")#      Result.Add(PadChar('Developer Name', TITLE_PAD_LENGTH, ' ') +  FReleaseUser);
Result.append('Project'.ljust(TITLE_PAD_LENGTH))#      Result.Add(PadChar('Project', TITLE_PAD_LENGTH, ' ')  +  Trim(aProjectsName));
Result.append('WDM Number'.ljust(TITLE_PAD_LENGTH))#      Result.Add(PadChar('WDM Number', TITLE_PAD_LENGTH, ' ')  +  FVersionNum);
Result.append('Tidzll Type'.ljust(TITLE_PAD_LENGTH))#      Result.Add(PadChar('Tidzll Type', TITLE_PAD_LENGTH, ' ')  +  Trim(aProjectsName));
      

Result.append('')#      Result.Add('');


Result.append('Details'.ljust(TITLE_PAD_LENGTH))#      Result.Add(PadChar('Details', 20, ' '));
Result.append('--------')#      Result.Add('--------');


Result.append('')#      Result.Add('');
#//      Result.Add('');
Result.append('QA Instructions:')#      Result.Add('QA Instructions:');
#//      Result.Add('');
#//      Result.Add('');
Result.append('----------------')#      Result.Add('----------------');
#//      Result.Add('');
#//      Result.Add('');
Result.append('Files Included:')#      Result.Add('Files Included:');
Result.append('----------------')#      Result.Add('----------------');


#//      Result.Add('');
Result.append('')#      Result.Add('');
Result.append('Issues Included:')#      Result.Add('Issues Included:');
Result.append('----------------')#      Result.Add('----------------');



f = open("test.txt",'w') 
Result = map(lambda x: x +'\n', Result)
f.writelines(Result)
f.close()
f = open("test.txt",'r')
print(f.read())
f.close()
